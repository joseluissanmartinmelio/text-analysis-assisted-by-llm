from docx import Document


def extract_text_without_footer_header(docx_path: str) -> str:
    doc = Document(docx_path)

    skip_text = set()
    for section in doc.sections:
        for para in section.header.paragraphs:
            txt = para.text.strip()
            if txt:
                skip_text.add(txt)
        for para in section.footer.paragraphs:
            txt = para.text.strip()
            if txt:
                skip_text.add(txt)

    body_text = []
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt and txt not in skip_text:
            body_text.append(txt)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    txt = para.text.strip()
                    if txt and txt not in skip_text:
                        body_text.append(txt)

    return "\n".join(body_text)
